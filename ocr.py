from datetime import datetime
import easyocr

# Initialize OCR once
reader = easyocr.Reader(['en'], gpu=False)


def extract_text(image_path):
    """
    Extract text from image using EasyOCR
    """

    results = reader.readtext(
        image_path,
        detail=0,
        paragraph=True
    )

    return "\n".join(results)


def save_as_txt(text, output_file):

    with open(output_file, "w", encoding="utf-8") as f:
        f.write(text)


def save_as_docx(text, output_file):

    from docx import Document

    doc = Document()

    doc.add_heading("OCR Output", level=1)

    doc.add_paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )

    doc.add_paragraph("")

    doc.add_paragraph(text)

    doc.save(output_file)


def save_as_pdf(text, output_file):

    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer
    )

    from reportlab.lib.styles import getSampleStyleSheet

    text = text.replace("&", "&amp;")
    text = text.replace("<", "&lt;")
    text = text.replace(">", "&gt;")
    text = text.replace("\n", "<br/>")

    doc = SimpleDocTemplate(output_file)

    styles = getSampleStyleSheet()

    story = []

    story.append(
        Paragraph(
            "OCR Output",
            styles["Title"]
        )
    )

    story.append(Spacer(1, 12))

    story.append(
        Paragraph(
            datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            styles["Normal"]
        )
    )

    story.append(Spacer(1, 12))

    story.append(
        Paragraph(
            text,
            styles["BodyText"]
        )
    )

    doc.build(story)